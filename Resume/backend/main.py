from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os, re, tempfile, shutil
from pathlib import Path

import pdfplumber
from docx import Document

app = FastAPI(title="Resume Scanner API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = Path("./outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


# Extracting Text

def extract_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()


def extract_docx(path: str) -> str:
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


# Section Detection

EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE    = re.compile(r"(\+?[\d][\d\s\-().]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[^\s,]+", re.I)
GITHUB_RE   = re.compile(r"github\.com/[^\s,]+", re.I)
URL_RE      = re.compile(r"https?://[^\s,]+", re.I)

SECTIONS = {
    "summary":        ["summary", "objective", "profile", "about me", "about", "overview", "career objective"],
    "experience":     ["experience", "work experience", "employment history", "work history", "professional experience"],
    "education":      ["education", "academic background", "qualifications", "schooling"],
    "skills":         ["skills", "technical skills", "core competencies", "technologies", "key skills", "competencies"],
    "projects":       ["projects", "personal projects", "key projects", "portfolio"],
    "certifications": ["certifications", "certificates", "licenses", "credentials"],
    "languages":      ["languages", "language proficiency", "spoken languages"],
    "interests":      ["interests", "hobbies", "activities", "extracurricular"],
    "references":     ["references", "referees"],
    "achievements":   ["achievements", "awards", "honors", "accomplishments"],
}

# Skill names

SKILL_CANONICAL: dict[str, str] = {
    "node.js": "Node.js",
    "node":    "Node.js",
    "react native": "React Native",
    "react":   "React",
    "vue":     "Vue.js",
    "svelte":  "Svelte",
    "angular": "Angular",
    "next.js": "Next.js",
    "express": "Express",
    "django":  "Django",
    "flask":   "Flask",
    "fastapi": "FastAPI",
    "spring":  "Spring",
    "laravel": "Laravel",
    "python":  "Python",
    "java":    "Java",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "kotlin":  "Kotlin",
    "swift":   "Swift",
    "ruby":    "Ruby",
    "php":     "PHP",
    "c++":     "C++",
    "c#":      "C#",
    ".net":    ".NET",
    "go":      "Go",
    "rust":    "Rust",
    "dart":    "Dart",
    "flutter": "Flutter",
    "sql":     "SQL",
    "mysql":   "MySQL",
    "postgresql": "PostgreSQL",
    "mongodb": "MongoDB",
    "redis":   "Redis",
    "sqlite":  "SQLite",
    "oracle":  "Oracle",
    "cassandra": "Cassandra",
    "elasticsearch": "Elasticsearch",
    "kafka":   "Kafka",
    "rabbitmq": "RabbitMQ",
    "docker":  "Docker",
    "kubernetes": "Kubernetes",
    "aws":     "AWS",
    "azure":   "Azure",
    "gcp":     "GCP",
    "terraform": "Terraform",
    "ansible": "Ansible",
    "jenkins": "Jenkins",
    "ci/cd":   "CI/CD",
    "git":     "Git",
    "github":  "GitHub",
    "gitlab":  "GitLab",
    "html":    "HTML",
    "css":     "CSS",
    "sass":    "Sass",
    "tailwind": "Tailwind CSS",
    "bootstrap": "Bootstrap",
    "rest":    "REST API",
    "graphql": "GraphQL",
    "grpc":    "gRPC",
    "linux":   "Linux",
    "bash":    "Bash",
    "shell":   "Shell",
    "tensorflow": "TensorFlow",
    "pytorch": "PyTorch",
    "scikit-learn": "Scikit-learn",
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "nlp":     "NLP",
    "data analysis": "Data Analysis",
    "pandas":  "Pandas",
    "numpy":   "NumPy",
    "excel":   "Excel",
    "tableau": "Tableau",
    "power bi": "Power BI",
    "figma":   "Figma",
    "photoshop": "Photoshop",
    "illustrator": "Illustrator",
    "agile":   "Agile",
    "scrum":   "Scrum",
    "jira":    "Jira",
    "firebase": "Firebase",
    "supabase": "Supabase",
    "nginx":   "Nginx",
}

# Skills considered to match with roles
REQUIRED_SKILLS = {
    "Python", "JavaScript", "TypeScript", "SQL", "Git",
    "Docker", "React", "Node.js", "AWS", "REST API",
    "Linux", "Agile", "HTML", "CSS",
}


def detect_sections(lines: list[str]) -> dict[str, int]:
    found: dict[str, int] = {}
    for i, line in enumerate(lines):
        norm = line.lower().strip().rstrip(":").strip()
        for sec, keywords in SECTIONS.items():
            if norm in keywords and sec not in found:
                found[sec] = i
    return found


def get_section_text(lines: list[str], smap: dict[str, int], key: str) -> str:
    if key not in smap:
        return ""
    start = smap[key] + 1
    nexts = sorted(v for k, v in smap.items() if k != key and v > start)
    end   = nexts[0] if nexts else len(lines)
    return "\n".join(lines[start:end]).strip()


# Personal Details Extraction 

def extract_name(lines: list[str]) -> str:
    for line in lines[:6]:
        s = line.strip()
        if (
            s and 2 <= len(s.split()) <= 5
            and not EMAIL_RE.search(s)
            and not PHONE_RE.search(s)
            and not any(s.lower().startswith(kw) for kws in SECTIONS.values() for kw in kws)
            and not any(c.isdigit() for c in s)
        ):
            return s
    return ""


# Skill Extraction & Normalization

def extract_skills(text: str) -> list[str]:

    # Matching skills from SKILL_CANONICAL against the text. Returns deduplicated, display-name list.
    
    lower = text.lower()
    found_display: set[str] = set()

    # Sort patterns longest-first to give multi-word patterns priority
    for pattern in sorted(SKILL_CANONICAL, key=len, reverse=True):
        display = SKILL_CANONICAL[pattern]
        if display in found_display:
            continue 
        if re.search(r"\b" + re.escape(pattern) + r"\b", lower):
            found_display.add(display)

    return sorted(found_display)


def rank_skills(skills: list[str]) -> list[str]:

    # Return skills sorted: commonly required first, then the rest.
    
    required = sorted(s for s in skills if s in REQUIRED_SKILLS)
    rest = sorted(s for s in skills if s not in REQUIRED_SKILLS)
    return required + rest

